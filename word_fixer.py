import os
import zipfile
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from lxml import etree

def get_page_count(docx_path):
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'docProps/app.xml' in z.namelist():
                xml_content = z.read('docProps/app.xml')
                tree = etree.fromstring(xml_content)
                ns = {'app': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                pages_elem = tree.find('app:Pages', ns)
                if pages_elem is not None and pages_elem.text:
                    return int(pages_elem.text)
    except Exception:
        pass
    return None

class SafeDocChecker:
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.messages = []

    def log(self, text, fixed=False):
        prefix = "[已修复] " if fixed else "[不符合] "
        self.messages.append(prefix + text)

    def is_in_table(self, paragraph):
        parent = paragraph._element.getparent()
        while parent is not None:
            if parent.tag == qn('w:tbl'):
                return True
            parent = parent.getparent()
        return False

    def check_and_fix_page_margins(self):
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
        self.log("所有节页边距已设置为：上2.5cm，其余2cm", fixed=True)

    def check_and_fix_headers_footers(self):
        for sectPr in self.doc.element.body.iter(qn('w:sectPr')):
            for tag in (qn('w:headerReference'), qn('w:footerReference')):
                for ref in sectPr.findall(tag):
                    sectPr.remove(ref)
                    self.log("已删除页眉/页脚/页码", fixed=True)

    def check_and_fix_paragraphs(self):
        for para in self.doc.paragraphs:
            if self.is_in_table(para):
                continue
            pf = para.paragraph_format
            if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or abs(pf.line_spacing - Pt(28)) > 1:
                pf.line_spacing = Pt(28)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                self.log(f"行距已设为固定值28磅（{para.text[:20]}...）", fixed=True)

            for run in para.runs:
                rPr = run._element.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(run._element, qn('w:rPr'))
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = etree.SubElement(rPr, qn('w:rFonts'))

                west = run.font.name
                if west and west != '宋体':
                    run.font.name = '宋体'
                    self.log("西文字体已改为宋体", fixed=True)
                if not west:
                    run.font.name = '宋体'

                east = rFonts.get(qn('w:eastAsia'))
                if east and east != '宋体':
                    rFonts.set(qn('w:eastAsia'), '宋体')
                    self.log("中文字体已改为宋体", fixed=True)
                if not east:
                    rFonts.set(qn('w:eastAsia'), '宋体')

                size = run.font.size
                if size and abs(size - Pt(14)) > 1:
                    run.font.size = Pt(14)
                    self.log("字号已改为四号(14磅)", fixed=True)
                if not size:
                    run.font.size = Pt(14)

                if run.font.bold:
                    run.font.bold = False
                    self.log("加粗已移除", fixed=True)
                if run.font.italic:
                    run.font.italic = False
                    self.log("倾斜已移除", fixed=True)
                if run.font.underline:
                    run.font.underline = False
                    self.log("下划线已移除", fixed=True)

                color = run.font.color.rgb
                if color and color != RGBColor(0, 0, 0):
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    self.log("文字颜色已改为黑色", fixed=True)
                if not color:
                    run.font.color.rgb = RGBColor(0, 0, 0)

        self.log("提示：图表、形状颜色请手动确认黑色", fixed=False)

    def check_toc(self):
        for instr in self.doc.element.body.iter(qn('w:instrText')):
            if instr.text and 'TOC' in instr.text:
                self.log("存在目录域(TOC)，请手动删除", fixed=False)
                return
        for para in self.doc.paragraphs:
            if para.style.name.startswith('TOC'):
                self.log("存在目录样式段落，疑似目录，请手动检查", fixed=False)
                break

    def run(self):
        self.check_toc()
        self.check_and_fix_page_margins()
        self.check_and_fix_headers_footers()
        self.check_and_fix_paragraphs()
        if not self.messages:
            self.messages.append("文档格式完全符合要求。")
        return self.messages

    def save(self, output_path):
        self.doc.save(output_path)


class App:
    def __init__(self, root):
        self.root = root
        root.title("Word格式检查与修复工具")
        root.geometry("750x600")
        self.filepath = tk.StringVar()

        frame1 = tk.Frame(root)
        frame1.pack(pady=10, padx=10, fill=tk.X)
        tk.Label(frame1, text="选择Word文件:").pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.filepath, width=50).pack(side=tk.LEFT, padx=5)
        tk.Button(frame1, text="浏览", command=self.browse).pack(side=tk.LEFT)

        frame2 = tk.Frame(root)
        frame2.pack(pady=5)
        tk.Button(frame2, text="检查并修复", command=self.check_and_fix,
                  bg="#4CAF50", fg="white", width=20, height=2).pack()

        frame3 = tk.Frame(root)
        frame3.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        tk.Label(frame3, text="检查结果与修复记录：").pack(anchor=tk.W)
        self.output_area = scrolledtext.ScrolledText(frame3, wrap=tk.WORD,
                                                     width=80, height=20, font=("宋体", 11))
        self.output_area.pack(fill=tk.BOTH, expand=True)

    def browse(self):
        filename = filedialog.askopenfilename(filetypes=[("Word文件", "*.docx")])
        if filename:
            self.filepath.set(filename)

    def check_and_fix(self):
        path = self.filepath.get()
        if not path:
            messagebox.showwarning("未选择文件", "请先选择一个Word文档")
            return
        if not os.path.exists(path) or not path.endswith('.docx'):
            messagebox.showerror("错误", "请选择有效的.docx文件")
            return

        dir_name = os.path.dirname(path)
        base = os.path.splitext(os.path.basename(path))[0]
        default_name = f"{base}_修正版.docx"
        save_path = filedialog.asksaveasfilename(
            title="保存修复后的文件",
            initialdir=dir_name,
            initialfile=default_name,
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx")]
        )
        if not save_path:
            return

        try:
            original_pages = get_page_count(path)
            checker = SafeDocChecker(path)
            messages = checker.run()
            checker.save(save_path)
            fixed_pages = get_page_count(save_path)
        except Exception as e:
            messagebox.showerror("处理失败", f"错误信息：{str(e)}")
            return

        orig_str = f"{original_pages} 页" if original_pages is not None else "无法统计"
        fixed_str = f"{fixed_pages} 页" if fixed_pages is not None else "无法统计"
        page_info = f"原文档页数：{orig_str}\n修复后文档页数：{fixed_str}"

        self.output_area.delete(1.0, tk.END)
        self.output_area.insert(tk.END, f"修复后文件已保存至：\n{save_path}\n\n{page_info}\n\n检查与修复记录：\n")
        for msg in messages:
            self.output_area.insert(tk.END, f"• {msg}\n")


if __name__ == "__main__":
    root = tk.Tk()
    App(root)
    root.mainloop()